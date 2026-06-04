import re
import sys
from pathlib import Path

from docx import Document
from docx.shared import Pt, Cm


DATA_DIR = Path(__file__).resolve().parent.parent.parent / "data"
OUTPUT_DIR = DATA_DIR / "output"
SKILLS_DIR = DATA_DIR / "skills"


sys.path.insert(0, str(SKILLS_DIR))


def _add_formatted_paragraph(doc, text: str, style: str = "Normal", bold: bool = False, italic: bool = False, size: int = 12):
    """Agrega un parrafo con formato inline: **bold** y *italic*."""
    p = doc.add_paragraph()
    pattern = r'(\*\*(.+?)\*\*|\*(.+?)\*|(.+?))(?=\*\*|\*|$)'
    pos = 0
    parts = []
    while pos < len(text):
        m = re.search(r'\*\*(.+?)\*\*|\*(.+?)\*', text[pos:])
        if not m:
            parts.append((text[pos:], False, False))
            break
        if m.start() > 0:
            parts.append((text[pos:pos+m.start()], False, False))
        if m.group(1):
            parts.append((m.group(1), True, False))
        else:
            parts.append((m.group(2), False, True))
        pos += m.end()
    if not parts:
        parts = [(text, False, False)]
    for txt, b, i in parts:
        if not txt:
            continue
        run = p.add_run(txt)
        run.bold = b or bold
        run.italic = i or italic
        run.font.size = Pt(size)
    return p


def _infer_title(document_content: str) -> str | None:
    first_line = document_content.strip().split("\n")[0] if document_content else ""
    title = first_line.replace("#", "").replace("*", "").replace('"', "").strip()
    if title:
        words = title.split()
        if len(words) > 20:
            title = " ".join(words[:20]) + "..."
        return title
    return None


def generate_docx(title: str, author: str, tutor: str, skill_type: str, document_content: str, logo_path: str | None = None, modelo: dict | None = None) -> Path:
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

    if logo_path and not Path(logo_path).exists():
        logo_path = None

    if not title or not title.strip():
        title = _infer_title(document_content) or "Documento sin título"
    
    if not author or not author.strip():
        author = "Sin asignar"
    
    if not tutor or not tutor.strip():
        tutor = "Sin asignar"

    skill_base = SKILLS_DIR / "base_documento.py"
    if skill_base.exists():
        import importlib.util
        spec = importlib.util.spec_from_file_location("base_documento", str(skill_base))
        mod = importlib.util.module_from_spec(spec)
        sys.modules["base_documento"] = mod
        spec.loader.exec_module(mod)

        doc = mod.crear_documento()
        mod.configurar_estilos(doc)

        section = doc.sections[0]
        mod.configurar_header(doc, section, fmt="decimal", start=1, mostrar=False)
        mod.configurar_primer_header_vacio(doc, section)

        mod.agregar_caratula(doc, title, author, tutor, logo_path=logo_path, modelo=modelo)

        doc.add_page_break()
        mod.agregar_toc(doc)
        doc.add_page_break()

        lines = document_content.split("\n")
        i = 0
        while i < len(lines):
            stripped = lines[i].strip()
            if not stripped:
                doc.add_paragraph("")
                i += 1
            elif stripped.startswith("|") and stripped.endswith("|") and i + 2 < len(lines) and "---" in lines[i+1]:
                table_lines = []
                while i < len(lines) and lines[i].strip().startswith("|"):
                    table_lines.append(lines[i].strip())
                    i += 1
                if len(table_lines) >= 3:
                    headers = [h.strip() for h in table_lines[0].strip("|").split("|")]
                    data = []
                    for row_line in table_lines[2:]:
                        cells = [c.strip() for c in row_line.strip("|").split("|")]
                        data.append(cells)
                    mod.add_table_with_data(doc, headers, data)
            elif stripped.startswith("### "):
                mod.add_heading_custom(doc, stripped[4:], level=3)
                i += 1
            elif stripped.startswith("## "):
                mod.add_heading_custom(doc, stripped[2:], level=2)
                i += 1
            elif stripped.startswith("# "):
                _add_formatted_paragraph(doc, stripped[1:].strip(), bold=True, size=14)
                i += 1
            elif stripped.startswith("- ") or stripped.startswith("* "):
                mod.add_bullets(doc, [stripped[2:]])
                i += 1
            else:
                mod.add_paragraph(doc, stripped)
                i += 1

        output_path = OUTPUT_DIR / f"{title.replace(' ', '_')[:50]}.docx"
        doc.save(str(output_path))
        return output_path

    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()
    for s in doc.sections:
        s.page_width = Cm(21.59)
        s.page_height = Cm(27.94)
        s.top_margin = Cm(2.54)
        s.bottom_margin = Cm(2.54)
        s.left_margin = Cm(3.0)
        s.right_margin = Cm(2.54)

    p = doc.add_paragraph()
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(16)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p = doc.add_paragraph()
    run = p.add_run(f"Autor: {author}  |  Tutor: {tutor}")
    run.font.size = Pt(12)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph("")
    doc.add_paragraph("")

    lines = document_content.split("\n")
    for line in lines:
        stripped = line.strip()
        if not stripped:
            doc.add_paragraph("")
        else:
            _add_formatted_paragraph(doc, stripped)

    output_path = OUTPUT_DIR / f"{title.replace(' ', '_')[:50]}.docx"
    doc.save(str(output_path))
    return output_path
