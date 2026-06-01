import os
import json
from pathlib import Path
from typing import Optional
from fastapi import APIRouter, HTTPException, UploadFile, File, Form, Header
from fastapi.responses import FileResponse
from app.models.schemas import GenerateRequest, GenerateDirectRequest
from app.services import file_reader, skill_loader, document_generator, gemini_service, web_search

DEFAULT_SKILL_ID = "indice"
router = APIRouter(prefix="/api/documents", tags=["documents"])


def _get_default_skill() -> dict | None:
    skill = skill_loader.get_skill(DEFAULT_SKILL_ID)
    if skill:
        return skill
    skills = skill_loader.list_skills()
    return skill_loader.get_skill(skills[0]["id"]) if skills else None


@router.post("/upload")
async def upload_file(file: UploadFile = File(...)):
    file_reader.ensure_upload_dir()
    upload_dir = Path(file_reader.UPLOAD_DIR)

    filepath = upload_dir / file.filename
    content = await file.read()
    with open(filepath, "wb") as f:
        f.write(content)

    info = file_reader.get_file_summary(str(filepath))
    return {
        "filename": info["filename"],
        "extension": info["extension"],
        "size_kb": info["size_kb"],
        "preview": info["preview"][:500],
    }


@router.get("/uploads")
def list_uploads():
    file_reader.ensure_upload_dir()
    upload_dir = Path(file_reader.UPLOAD_DIR)
    files = []
    for f in sorted(upload_dir.iterdir()):
        if f.is_file():
            info = file_reader.get_file_summary(str(f))
            files.append({
                "filename": info["filename"],
                "extension": info["extension"],
                "size_kb": info["size_kb"],
                "summary": info["summary"],
            })
    return files


@router.delete("/uploads/{filename}")
def delete_upload(filename: str):
    upload_dir = Path(file_reader.UPLOAD_DIR)
    filepath = upload_dir / filename
    if filepath.exists() and filepath.is_file():
        filepath.unlink()
        return {"status": "deleted", "filename": filename}
    raise HTTPException(status_code=404, detail="Archivo no encontrado")


@router.post("/generate")
def generate_document(data: GenerateRequest, x_gemini_key: Optional[str] = Header(None)):
    skill = skill_loader.get_skill(data.skill_id)
    if not skill:
        raise HTTPException(status_code=404, detail="Skill no encontrada")

    uploaded_content = file_reader.extract_all_uploaded_content()

    document_content = gemini_service.generate_document_content(
        skill_content=skill["content"],
        uploaded_content=uploaded_content,
        user_input=data.content_input,
        api_key=x_gemini_key,
    )

    logo_path = "/home/eynor/Documentos/Biblioteca/logoUnifranz/logounifranz.png"

    output_path = document_generator.generate_docx(
        title=data.title,
        author=data.author,
        tutor=data.tutor,
        skill_type=skill["type"],
        document_content=document_content,
        logo_path=logo_path if os.path.exists(logo_path) else None,
    )

    return {
        "status": "success",
        "download_url": f"/api/documents/download/{output_path.name}",
        "filename": output_path.name,
    }


@router.get("/download/{filename}")
def download_document(filename: str):
    output_dir = Path(document_generator.OUTPUT_DIR)
    filepath = output_dir / filename
    if not filepath.exists():
        raise HTTPException(status_code=404, detail="Archivo no encontrado")
    return FileResponse(
        path=str(filepath),
        filename=filename,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )


@router.post("/generate-direct")
def generate_direct(data: GenerateDirectRequest):
    topic = data.topic or data.title
    results = web_search.search_web(topic, max_results=8)
    if results:
        lines = [f"# {topic}", ""]
        for i, r in enumerate(results, 1):
            lines.append(f"## {i}. {r['title']}")
            lines.append("")
            lines.append(r['body'])
            lines.append("")
        document_content = "\n".join(lines)
    else:
        document_content = f"# {topic}\n\nDocumento generado sobre: {topic}"

    default_skill = _get_default_skill()
    skill_type = default_skill["type"] if default_skill else "md"

    logo_path = "/home/eynor/Documentos/Biblioteca/logoUnifranz/logounifranz.png"

    output_path = document_generator.generate_docx(
        title=data.title,
        author=data.author or "",
        tutor=data.tutor or "",
        skill_type=skill_type,
        document_content=document_content,
        logo_path=logo_path if os.path.exists(logo_path) else None,
    )

    return {
        "status": "success",
        "download_url": f"/api/documents/download/{output_path.name}",
        "filename": output_path.name,
    }


@router.post("/generate-demo")
def generate_demo():
    default_skill = _get_default_skill()

    logo_path = "/home/eynor/Documentos/Biblioteca/logoUnifranz/logounifranz.png"

    demo_content = """# Documento de Prueba DocGent

## 1. Introduccion

Este documento es una **demostracion** de todas las capacidades de generacion de DocGent. Incluye *texto en cursiva*, **texto en negrita**, y combinaciones de **formato *mixto***.

## 2. Listas y viñetas

- **Distribuciones Linux**: Ubuntu, Fedora, Debian
- **Lenguajes de programacion**: Python, JavaScript, Rust
- **Base de datos**: PostgreSQL, MySQL, MongoDB

## 3. Tabla comparativa

| Distribucion | Base | Gestor Paquetes | Enfoque |
|---|---|---|---|
| Ubuntu | Debian | APT | Escritorio / Servidor |
| Fedora | Red Hat | DNF | Tecnologia reciente |
| Debian | Independiente | APT | Estabilidad |
| Arch | Independiente | Pacman | Rolling release |
| openSUSE | Independiente | Zypper | Desarrollo / Sysadmin |

## 4. Contenido academico

El siguiente analisis presenta una **comparacion detallada** de los sistemas operativos modernos.

### 4.1 Arquitectura del kernel

El *kernel* de Linux es un nucleo monolitico que soporta:

1. Gestion de procesos
2. Gestion de memoria
3. Sistema de archivos virtual
4. Controladores de dispositivos
5. Redes y conectividad

### 4.2 Formula matematica

La eficiencia de un sistema operativo puede medirse como:

E = (T_procesamiento + T_entrada_salida) / T_total

Donde **E** representa la eficiencia y los valores se miden en milisegundos.

## 5. Datos estructurados

| Caracteristica | Linux | Windows | macOS |
|---|---|---|---|
| Codigo fuente | Abierto | Cerrado | Hibrido |
| Costo | Gratuito | Licencia paga | Gratuito (hardware Apple) |
| Personalizacion | Total | Limitada | Parcial |
| Seguridad | Alta | Media | Alta |
| Rendimiento servidor | Excelente | Bueno | Limitado |

## 6. Referencias

- **Tanenbaum, A. S.** (2015). *Modern Operating Systems*. Pearson.
- **Silberschatz, A.** (2018). *Operating System Concepts*. Wiley.
- Documentacion oficial de **Linux Kernel** (2025). kernel.org
- **Rago, S.** (2023). *The Design of the UNIX Operating System*. Prentice Hall.

## 7. Conclusion

DocGent permite la generacion automatica de documentos .docx con **formato APA 7**, incluyendo:

- Caratula institucional con logo
- Tabla de contenidos automatica
- Encabezados estilizados (Heading 2 y 3)
- Parrafos justificados con sangria
- **Tablas** con bordes y formato APA
- *Listas* con viñetas y numeradas
- Numeracion de pagina
- Referencias bibliograficas
- **Negritas** y *cursivas* en linea
"""

    output_path = document_generator.generate_docx(
        title="Documento de Prueba DocGent",
        author="DocGent Demo",
        tutor="Sistema de Pruebas",
        skill_type="md",
        document_content=demo_content,
        logo_path=logo_path if os.path.exists(logo_path) else None,
    )

    return {
        "status": "success",
        "download_url": f"/api/documents/download/{output_path.name}",
        "filename": output_path.name,
    }


@router.get("/list")
def list_documents():
    output_dir = Path(document_generator.OUTPUT_DIR)
    output_dir.mkdir(parents=True, exist_ok=True)
    files = []
    for f in sorted(output_dir.iterdir()):
        if f.suffix == ".docx":
            files.append({
                "filename": f.name,
                "size_kb": round(f.stat().st_size / 1024, 2),
                "created": f.stat().st_mtime,
            })
    return files


@router.get("/preview/{filename}")
def preview_document(filename: str):
    from docx import Document as DocxDocument
    output_dir = Path(document_generator.OUTPUT_DIR)
    filepath = output_dir / filename
    if not filepath.exists():
        raise HTTPException(status_code=404, detail="Archivo no encontrado")
    doc = DocxDocument(str(filepath))
    paragraphs = []
    for p in doc.paragraphs:
        runs = []
        for r in p.runs:
            runs.append({
                "text": r.text,
                "bold": r.bold,
                "italic": r.italic,
                "size": r.font.size,
            })
        paragraphs.append({
            "text": p.text,
            "style": p.style.name if p.style else "Normal",
            "runs": runs,
        })
    tables = []
    for t in doc.tables:
        rows = []
        for row in t.rows:
            cells = [cell.text for cell in row.cells]
            rows.append(cells)
        tables.append(rows)
    return {
        "filename": filename,
        "paragraphs": paragraphs,
        "tables": tables,
        "size_kb": round(filepath.stat().st_size / 1024, 2),
    }
